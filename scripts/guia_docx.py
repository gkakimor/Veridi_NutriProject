"""
Gera os guias passo a passo em Word, um por TRILHA.

Por que um documento por trilha, e não um manual único: quem vai criar um
produto pelo Projeto não quer folhear o caminho direto para achar o seu. Um
guia que serve a todos os caminhos obriga cada leitor a ignorar metade dele,
e é assim que manual vira enfeite.

As etapas — o que fazer, o que esperar, a regra e o que trava — vêm de
`handoff/guia/manifest.json`, produzido por `scripts/guia-capturas.mjs`, que
percorre o sistema de verdade e captura uma imagem por passo. A moldura de
cada documento (para quem é, o que precisa existir antes, o vocabulário, o
que fica pronto no fim) é escrita aqui.

    python scripts/guia_docx.py

Saída: um `.docx` por trilha em `docs/`. `docs/*.docx` é ignorado pelo git —
são artefatos grandes e regeráveis, não fonte.
"""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
MANIFESTO = RAIZ / "handoff" / "guia" / "manifest.json"
SAIDA = RAIZ / "docs"

# Identidade Veridi, os mesmos valores de `tokens.css`. Repetir aqui é
# deliberado: o Word não lê CSS, e um verde "quase igual" salta aos olhos
# quando o documento fica ao lado da tela.
VERDE_ESCURO = RGBColor(0x0C, 0x36, 0x29)
VERDE = RGBColor(0x1B, 0x5E, 0x43)
TINTA = RGBColor(0x17, 0x25, 0x1E)
TINTA_FRACA = RGBColor(0x5A, 0x6B, 0x63)
ALERTA = RGBColor(0x8A, 0x5A, 0x1C)

LARGURA_IMAGEM = Inches(6.3)

# ---------------------------------------------------------------------------
# O vocabulário. Igual nas duas trilhas porque é o mesmo sistema — e é o que
# separa "sequência de cliques" de "guia": quem não sabe o que é versão ativa
# executa o passo e não entende por que ele existe.
# ---------------------------------------------------------------------------

CONCEITOS = [
    ("Projeto",
     "A negociação, que existe ANTES do produto. Guarda o briefing do cliente, "
     "as amostras e as propostas. Aprovar o projeto é o que faz o produto nascer."),
    ("Produto",
     "O que a Veridi fabrica e vende: nome, cliente, perfil industrial, "
     "formulação, custo e preço. Não é o registro de estoque — todo produto tem "
     "um item de produto acabado ligado a ele, criado junto, e é esse item que "
     "tem lote, validade e saldo."),
    ("Formulação e versão ativa",
     "A receita oficial do produto. É escrita em versões: enquanto é rascunho "
     "muda à vontade, e vira documento no momento em que é ativada. Só existe "
     "uma versão ativa por produto, e é ela que a produção executa e que o custo "
     "e o preço leem."),
    ("Estrutura de custos",
     "Declara sobre qual versão da receita se calcula, sobre qual base de "
     "produção, e quais recursos, energia e premissas entram. Também é escrita "
     "em versões, e também precisa ser ativada."),
    ("Cálculo",
     "Aplica as referências de custo de uma data à estrutura e CONGELA o "
     "resultado num documento com código próprio. É ele que a precificação lê — "
     "nunca um custo recalculado na hora da venda."),
    ("CMV",
     "O custo da mercadoria vendida do produto, montado a partir do cálculo "
     "congelado. Precisa de três coisas prontas: formulação ativa, estrutura "
     "ativa e cálculo salvo."),
]

PRE_REQUISITOS_COMUNS = [
    "As matérias-primas e embalagens da receita já cadastradas em Cadastros › "
    "Itens de estoque — e, para o custo fechar completo, já recebidas com preço, "
    "porque é do recebimento que vem a referência de custo.",
    "Os recursos industriais da estrutura (mão de obra, equipamento, energia) "
    "cadastrados em Gestão › Recursos industriais, cada um com TARIFA VIGENTE. "
    "Recurso sem tarifa trava a ativação da estrutura de custos.",
    "No cadastro do produto, dois campos que parecem opcionais e cobram depois: "
    "“Lote mínimo”, que vira a base de produção sugerida, e “Unidades por caixa”, "
    "que habilita premissa por caixa de expedição.",
]

NARRATIVA = {
    "trilhaA": {
        "arquivo": "Guia_Produto_e_CMV_via_Projeto.docx",
        "titulo": "Do projeto ao CMV",
        "subtitulo": "Criar um produto pelo funil comercial e chegar ao custo",
        "resumo": "Para quem recebe o briefing de um cliente, negocia a proposta e\n"
                  "precisa saber quanto o produto custa antes de precificar.",
        "antesDeComecar": [
            "O cliente já cadastrado. O projeto pertence a um cliente, e o botão "
            "de criar só habilita depois que ele é escolhido — se não existir, dá "
            "para cadastrar pelo próprio campo de busca, sem perder o formulário.",
        ] + PRE_REQUISITOS_COMUNS,
        "aoFim": [
            "Projeto aprovado, com a proposta aceita registrada.",
            "Produto criado e promovido a operacional, com item de produto acabado "
            "próprio.",
            "Formulação com versão ativa.",
            "Estrutura de custos ativa, sem pendências bloqueantes.",
            "Cálculo salvo — um documento congelado, com código próprio.",
            "CMV disponível, com o custo por lote e por unidade.",
        ],
    },
    "trilhaB": {
        "arquivo": "Guia_Produto_e_CMV_Direto.docx",
        "titulo": "Do cadastro ao CMV",
        "subtitulo": "Criar um produto direto pela tela e chegar ao custo",
        "resumo": "Para o produto que não passa pelo funil comercial — linha própria,\n"
                  "reformulação, ou um item que já vem decidido.",
        "antesDeComecar": [
            "O cliente já cadastrado. Cliente é obrigatório na criação do produto "
            "e trava depois que existir pedido, ordem de produção ou orçamento — "
            "se não existir, dá para cadastrar pelo próprio campo de busca.",
        ] + PRE_REQUISITOS_COMUNS,
        "aoFim": [
            "Produto criado já como operacional, com item de produto acabado "
            "próprio.",
            "Formulação com versão ativa.",
            "Estrutura de custos ativa, sem pendências bloqueantes.",
            "Cálculo salvo — um documento congelado, com código próprio.",
            "CMV disponível, com o custo por lote e por unidade.",
        ],
    },
}


def paragrafo(doc, texto, *, tamanho=11, cor=TINTA, negrito=False, italico=False,
              espaco_antes=0, espaco_depois=6, alinhamento=None):
    p = doc.add_paragraph()
    if alinhamento is not None:
        p.alignment = alinhamento
    run = p.add_run(texto)
    run.font.size = Pt(tamanho)
    run.font.color.rgb = cor
    run.bold = negrito
    run.italic = italico
    p.paragraph_format.space_before = Pt(espaco_antes)
    p.paragraph_format.space_after = Pt(espaco_depois)
    return p


def rotulo_e_texto(doc, rotulo, texto, *, cor_rotulo=VERDE):
    """Uma linha "Rótulo — texto", com o rótulo destacado.

    Cabeçalho próprio para cada campo gastaria meia página por etapa e
    empurraria a imagem seguinte para longe do texto que a explica.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{rotulo}  ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = cor_rotulo
    t = p.add_run(texto)
    t.font.size = Pt(10.5)
    t.font.color.rgb = TINTA
    return p


def capa(doc, fluxo, total_etapas):
    for _ in range(4):
        doc.add_paragraph()
    paragrafo(doc, "VERIDI NUTRITION", tamanho=11, cor=VERDE, negrito=True,
              alinhamento=WD_ALIGN_PARAGRAPH.CENTER, espaco_depois=2)
    paragrafo(doc, fluxo["titulo"], tamanho=26, cor=VERDE_ESCURO, negrito=True,
              alinhamento=WD_ALIGN_PARAGRAPH.CENTER, espaco_depois=4)
    paragrafo(doc, fluxo["subtitulo"], tamanho=13, cor=TINTA_FRACA,
              alinhamento=WD_ALIGN_PARAGRAPH.CENTER, espaco_depois=24)
    paragrafo(doc, fluxo["resumo"], tamanho=11, cor=TINTA,
              alinhamento=WD_ALIGN_PARAGRAPH.CENTER, espaco_depois=30)
    paragrafo(doc, f"{total_etapas} etapas  ·  " +
              date.today().strftime("gerado em %d/%m/%Y"), tamanho=9,
              cor=TINTA_FRACA, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_section(WD_SECTION.NEW_PAGE)


def secao(doc, titulo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = VERDE_ESCURO


def lista(doc, itens):
    for item in itens:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = TINTA


def etapa(doc, numero, passo):
    """Uma etapa: título numerado, o que fazer, a imagem, e as ressalvas.

    A ordem é deliberada. A pessoa lê o que tem de fazer, faz, e só então
    confere na imagem — imagem antes da instrução vira decoração.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{numero}.  {passo['titulo']}")
    r.bold = True
    r.font.size = Pt(13.5)
    r.font.color.rgb = VERDE_ESCURO

    if passo.get("oQueFazer"):
        rotulo_e_texto(doc, "O QUE FAZER", passo["oQueFazer"])
    if passo.get("oQueEsperar"):
        rotulo_e_texto(doc, "O QUE ESPERAR", passo["oQueEsperar"])

    caminho = RAIZ / passo["imagem"]
    if caminho.exists():
        img = doc.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.paragraph_format.space_before = Pt(9)
        img.paragraph_format.space_after = Pt(6)
        img.add_run().add_picture(str(caminho), width=LARGURA_IMAGEM)
    else:
        # Melhor dizer que a imagem falta do que entregar um guia com um
        # buraco silencioso onde deveria haver a prova do passo.
        paragrafo(doc, f"[imagem não encontrada: {passo['imagem']}]",
                  tamanho=9, cor=TINTA_FRACA, italico=True)

    if passo.get("regra"):
        rotulo_e_texto(doc, "POR QUÊ", passo["regra"])
    if passo.get("sePrender"):
        # Cor diferente: é a única parte que a pessoa procura DEPOIS de
        # travar, e procurar exige achar de relance.
        rotulo_e_texto(doc, "SE PRENDER", passo["sePrender"], cor_rotulo=ALERTA)


def documento(fluxo, passos, convergencia, saida: Path):
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    capa(doc, fluxo, len(passos))

    secao(doc, "Antes de começar")
    paragrafo(doc, "Nada aqui é cadastrado durante o caminho. Faltando um destes, "
                   "o processo trava no meio — e trava em telas diferentes da que "
                   "tem o problema.", tamanho=10.5, cor=TINTA_FRACA)
    lista(doc, fluxo["antesDeComecar"])

    secao(doc, "Os termos que aparecem no caminho")
    for termo, texto in CONCEITOS:
        rotulo_e_texto(doc, termo, texto)

    if convergencia:
        secao(doc, "Onde este caminho encontra o outro")
        paragrafo(doc, convergencia, tamanho=10.5)

    secao(doc, "Passo a passo")
    for i, passo in enumerate(passos, start=1):
        etapa(doc, i, passo)

    secao(doc, "O que fica pronto ao fim")
    lista(doc, fluxo["aoFim"])

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)
    return saida


def main():
    if not MANIFESTO.exists():
        sys.exit(f"Manifesto não encontrado: {MANIFESTO}")

    dados = json.loads(MANIFESTO.read_text(encoding="utf-8"))
    convergencia = dados.get("convergencia")
    if isinstance(convergencia, dict):
        convergencia = convergencia.get("texto") or json.dumps(convergencia, ensure_ascii=False)

    gerados = []
    for chave, fluxo in NARRATIVA.items():
        passos = dados.get(chave) or []
        if not passos:
            sys.exit(f"Trilha sem etapas no manifesto: {chave}")
        gerados.append(documento(fluxo, passos, convergencia, SAIDA / fluxo["arquivo"]))

    for c in gerados:
        print(f"{c.relative_to(RAIZ)}  ({c.stat().st_size / 1024 / 1024:.1f} MB)")


if __name__ == "__main__":
    main()
